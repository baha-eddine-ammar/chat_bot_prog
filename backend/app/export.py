from io import BytesIO
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document


def build_pdf_bytes(text: str) -> bytes:
	buffer = BytesIO()
	c = canvas.Canvas(buffer, pagesize=LETTER)
	width, height = LETTER
	left_margin = 1 * inch
	top = height - 1 * inch
	max_width = width - 2 * inch

	# Simple line wrapping
	from reportlab.pdfbase.pdfmetrics import stringWidth
	lines = []
	for paragraph in text.split("\n"):
		words = paragraph.split(" ")
		line = ""
		for w in words:
			candidate = (line + " " + w).strip()
			if stringWidth(candidate, "Helvetica", 11) <= max_width:
				line = candidate
			else:
				lines.append(line)
				line = w
		if line:
			lines.append(line)
		lines.append("")

	c.setFont("Helvetica", 11)
	y = top
	for line in lines:
		if y < 1 * inch:
			c.showPage()
			c.setFont("Helvetica", 11)
			y = top
		c.drawString(left_margin, y, line)
		y -= 14

	c.save()
	buffer.seek(0)
	return buffer.read()


def build_docx_bytes(text: str) -> bytes:
	doc = Document()
	for paragraph in text.split("\n\n"):
		p = doc.add_paragraph()
		for line in paragraph.split("\n"):
			p.add_run(line)
			p.add_run("\n")
	bio = BytesIO()
	doc.save(bio)
	bio.seek(0)
	return bio.read()
